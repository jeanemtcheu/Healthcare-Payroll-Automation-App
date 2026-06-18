import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


# =========================================================
# GENERATE DOCUMENTS FUNCTION
# =========================================================

def generate_documents(
    final_shift_hours_df,
    caretakers_file,
    output_folder="outputs"
):

    # =====================================================
    # CREATE PAY PERIOD LABEL
    # =====================================================

    df = final_shift_hours_df.copy()

    df["service_date"] = pd.to_datetime(df["service_date"])

    period_start = (
        df["service_date"]
        .min()
        .strftime("%Y-%m-%d")
    )

    period_end = (
        df["service_date"]
        .max()
        .strftime("%Y-%m-%d")
    )

    pay_period_label = (
        f"{period_start}_to_{period_end}"
    )

    # =====================================================
    # CREATE OUTPUT FOLDERS
    # =====================================================

    paystub_folder = os.path.join(
        output_folder,
        f"paystubs_{pay_period_label}"
    )

    timesheet_folder = os.path.join(
        output_folder,
        f"timesheets_{pay_period_label}"
    )

    os.makedirs(paystub_folder, exist_ok=True)
    os.makedirs(timesheet_folder, exist_ok=True)

    # =====================================================
    # CREATE PAYROLL SUMMARY
    # =====================================================

    payroll_summary = (
        df
        .groupby(
            ["caretaker_id", "caretaker_name"],
            as_index=False
        )
        .agg({
            "approved_shift_hours": "sum",
            "unpaid_over_cap_hours": "sum"
        })
    )

    payroll_summary = payroll_summary.rename(columns={
        "approved_shift_hours": "total_paid_hours",
        "unpaid_over_cap_hours": "total_unpaid_hours"
    })

    print("\nPAYROLL SUMMARY")
    print("-" * 50)
    print(payroll_summary)

    # =====================================================
    # LOAD CARETAKER PAY RATES
    # =====================================================

    caretakers_df = pd.read_csv(caretakers_file)

    payroll_summary = payroll_summary.merge(
        caretakers_df,
        on=["caretaker_id", "caretaker_name"],
        how="left"
    )

    # =====================================================
    # GROSS PAY
    # =====================================================

    payroll_summary["gross_pay"] = (
        payroll_summary["total_paid_hours"]
        * payroll_summary["hourly_rate"]
    ).round(2)

    # =====================================================
    # TAX DEDUCTIONS
    # =====================================================

    payroll_summary["social_security"] = (
        payroll_summary["gross_pay"] * 0.063
    ).round(2)

    payroll_summary["medicare"] = (
        payroll_summary["gross_pay"] * 0.0145
    ).round(2)

    payroll_summary["estimated_federal_tax"] = (
        payroll_summary["gross_pay"] * 0.124
    ).round(2)

    payroll_summary["md_state_tax"] = (
            payroll_summary["gross_pay"] * 0.124
    ).round(2)

    payroll_summary["total_deductions"] = (
        payroll_summary["social_security"]
        + payroll_summary["medicare"]
        + payroll_summary["estimated_federal_tax"]
        + payroll_summary["md_state_tax"]
    ).round(2)

    payroll_summary["net_pay"] = (
        payroll_summary["gross_pay"]
        - payroll_summary["total_deductions"]
    ).round(2)

    print("\nFINAL PAYROLL SUMMARY")
    print("-" * 50)
    print(payroll_summary)

    # =====================================================
    # CREATE STYLIZED TIMESHEETS
    # =====================================================

    print("\nCREATING STYLIZED TIMESHEETS")
    print("-" * 50)

    for caretaker_id, caretaker_data in df.groupby("caretaker_id"):

        caretaker_name = (
            caretaker_data["caretaker_name"]
            .iloc[0]
        )

        doc = Document()

        # -------------------------------------------------
        # HEADER
        # -------------------------------------------------

        title = doc.add_heading(
            "BYZANTINE TIMESHEET",
            level=1
        )

        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        company_info = doc.add_paragraph()

        company_info.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        company_info.add_run(
            "Byzantine Healthcare Services LLC\n"
            "12708 Willow Marsh Ln\n"
            "Bowie, MD 20720-4655\n"
            "Phone: (301)-906-5382"
        )

        doc.add_paragraph()

        doc.add_paragraph(
            f"Pay Period: "
            f"{period_start} to {period_end}"
        )

        # -------------------------------------------------
        # EMPLOYEE INFO
        # -------------------------------------------------

        doc.add_heading(
            "Caretaker Information",
            level=2
        )

        doc.add_paragraph(
            f"Employee Name: {caretaker_name}"
        )

        doc.add_paragraph(
            f"Caretaker ID: {caretaker_id}"
        )

        # -------------------------------------------------
        # SHIFT TABLE
        # -------------------------------------------------

        doc.add_heading(
            "Shift Details",
            level=2
        )

        table = doc.add_table(
            rows=1,
            cols=6
        )

        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells

        hdr_cells[0].text = "Service Date"
        hdr_cells[1].text = "Client"
        hdr_cells[2].text = "Clock In"
        hdr_cells[3].text = "Clock Out"
        hdr_cells[4].text = "Approved Hours"
        hdr_cells[5].text = "Over Cap Hours"

        for _, shift in caretaker_data.iterrows():

            row_cells = (
                table.add_row().cells
            )

            row_cells[0].text = (
                str(shift["service_date"])
            )

            row_cells[1].text = (
                str(shift["client_name"])
            )

            row_cells[2].text = (
                str(shift["corrected_clock_in"])
            )

            row_cells[3].text = (
                str(shift["corrected_clock_out"])
            )

            row_cells[4].text = (
                str(
                    round(
                        shift["approved_shift_hours"],
                        2
                    )
                )
            )

            row_cells[5].text = (
                str(
                    round(
                        shift["unpaid_over_cap_hours"],
                        2
                    )
                )
            )

        # -------------------------------------------------
        # SAVE TIMESHEET
        # -------------------------------------------------

        filename = os.path.join(
            timesheet_folder,
            (
                f"timesheet_"
                f"{caretaker_name.replace(' ', '_')}_"
                f"{pay_period_label}.docx"
            )
        )

        doc.save(filename)

        print(f"Created {filename}")

    # =====================================================
    # CREATE STYLIZED PAYSTUBS
    # =====================================================

    print("\nCREATING STYLIZED PAYSTUBS")
    print("-" * 50)

    for _, row in payroll_summary.iterrows():

        caretaker_name = row["caretaker_name"]

        doc = Document()

        # -------------------------------------------------
        # HEADER
        # -------------------------------------------------

        title = doc.add_heading(
            "EMPLOYEE PAYSTUB",
            level=1
        )

        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        company_info = doc.add_paragraph()

        company_info.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        company_info.add_run(
            "Byzantine Healthcare Services LLC\n"
            "Payroll Department\n"
            "12708 Willow Marsh Ln\n"
            "Bowie, MD 20720-4655"
        )

        doc.add_paragraph()

        doc.add_paragraph(
            f"Pay Period: "
            f"{period_start} to {period_end}"
        )

        # -------------------------------------------------
        # EMPLOYEE INFO
        # -------------------------------------------------

        doc.add_heading(
            "Employee Information",
            level=2
        )

        doc.add_paragraph(
            f"Employee Name: {caretaker_name}"
        )

        doc.add_paragraph(
            f"Caretaker ID: "
            f"{row['caretaker_id']}"
        )

        # -------------------------------------------------
        # PAYROLL TABLE
        # -------------------------------------------------

        doc.add_heading(
            "Payroll Summary",
            level=2
        )

        payroll_table = doc.add_table(
            rows=1,
            cols=2
        )

        payroll_table.style = "Table Grid"

        hdr_cells = payroll_table.rows[0].cells

        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Amount"

        payroll_items = [
            (
                "Total Paid Hours",
                row["total_paid_hours"]
            ),
            (
                "Hourly Rate",
                f"${row['hourly_rate']}"
            ),
            (
                "Gross Pay",
                f"${row['gross_pay']}"
            ),
            (
                "Social Security",
                f"${row['social_security']}"
            ),
            (
                "Medicare",
                f"${row['medicare']}"
            ),
            (
                "Federal Tax Estimate",
                f"${row['estimated_federal_tax']}"
            ),

            (
                "Total Paid Hours",
                row["total_paid_hours"]
            ),
            
            (
                "Total Deductions",
                f"${row['total_deductions']}"
            ),
            (
                "Net Pay",
                f"${row['net_pay']}"
            )
        ]

        for label, value in payroll_items:

            cells = (
                payroll_table.add_row().cells
            )

            cells[0].text = str(label)
            cells[1].text = str(value)

        # -------------------------------------------------
        # FINE PRINT
        # -------------------------------------------------

        doc.add_paragraph()

        fine_print = doc.add_paragraph()

        fine_print.add_run(
            "This paystub is confidential "
            "and intended solely for the "
            "employee listed above."
        ).italic = True

        # -------------------------------------------------
        # SAVE PAYSTUB
        # -------------------------------------------------

        filename = os.path.join(
            paystub_folder,
            (
                f"paystub_"
                f"{caretaker_name.replace(' ', '_')}_"
                f"{pay_period_label}.docx"
            )
        )

        doc.save(filename)

        print(f"Created {filename}")

    print("\nALL DOCUMENTS GENERATED SUCCESSFULLY")
    print("-" * 50)